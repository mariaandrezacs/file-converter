import io

from docx import Document

from src.word_to_pdf.converters.word_to_pdf import WordToPdfConverter


def test_convert_valid_word_to_pdf():
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("Ana")
    document.add_paragraph("João")
    document.save(buffer)
    buffer.seek(0)

    converter = WordToPdfConverter()
    result = converter.convert(buffer)

    assert result["rows_processed"] == 2
